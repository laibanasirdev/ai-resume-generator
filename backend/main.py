from fastapi import FastAPI, Form, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from groq import Groq, RateLimitError, APIStatusError
from dotenv import load_dotenv
from fpdf import FPDF
import PyPDF2
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import io
import os
import json

load_dotenv()

client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# Models tried in order — falls back automatically on rate-limit errors
_MODELS = [
    "llama-3.3-70b-versatile",   # primary:  100K TPD
    "llama-3.1-8b-instant",       # fallback: 500K TPD
    "gemma2-9b-it",               # last resort
]


def _chat(prompt: str, temperature: float = 0.4) -> str:
    """Call Groq with automatic model fallback on rate-limit or request-too-large errors."""
    last_err: Exception | None = None
    for model in _MODELS:
        try:
            resp = client.chat.completions.create(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=4000,
                temperature=temperature,
            )
            return resp.choices[0].message.content
        except RateLimitError as e:
            last_err = e
            continue
        except APIStatusError as e:
            # 413 = prompt too large for this model — try the next one
            if e.status_code == 413:
                last_err = e
                continue
            raise  # re-raise other API errors (auth, bad request, etc.)
    raise HTTPException(
        status_code=429,
        detail=f"All models are unavailable (rate-limited or prompt too large). ({last_err})",
    )

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── colours ────────────────────────────────────────────────────────────────────

_NAVY   = (22,  43,  77)    # header band background
_LIGHT  = (185, 200, 220)   # contact text in header
_ACCENT = (55,  48,  163)   # section heading colour
_DARK   = (17,  24,  39)    # body heading text
_MID    = (55,  65,  81)    # body / bullets
_RULE   = (210, 214, 230)   # section underline


# ── helpers ────────────────────────────────────────────────────────────────────

def split_content(full_text: str) -> tuple[str, str]:
    """Return (resume_section, cover_letter_section)."""
    pattern = re.compile(r'(?m)^#{1,3}\s*COVER\s*LETTER\s*$', re.IGNORECASE)
    match = pattern.search(full_text)
    if match:
        return full_text[: match.start()].strip(), full_text[match.start():].strip()
    return full_text.strip(), ""


_UNICODE_MAP = str.maketrans({
    '\u2013': '-',   '\u2014': '--',
    '\u2018': "'",   '\u2019': "'",
    '\u201c': '"',   '\u201d': '"',
    '\u2022': '-',   '\u25cf': '-',
    '\u00a0': ' ',   '\u2026': '...',
})


def _clean(text: str) -> str:
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',     r'\1', text)
    text = re.sub(r'__(.+?)__',     r'\1', text)
    text = re.sub(r'`(.+?)`',       r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]',  r'\1', text)
    text = text.translate(_UNICODE_MAP)
    return text.encode('latin-1', errors='ignore').decode('latin-1').strip()


# ── post-processing pipeline ───────────────────────────────────────────────────

_THINKING_LINE_RE = re.compile(
    r'^(INTERNAL\s+STEP\s*\d|STEP\s*[123]\s*[—\-]|━{10,})',
    re.IGNORECASE,
)


def _strip_thinking_lines(text: str) -> str:
    """Remove lines that are clearly leaked chain-of-thought (STEP headers, separator bars)."""
    return '\n'.join(
        ln for ln in text.split('\n')
        if not _THINKING_LINE_RE.match(ln.strip())
    )


def _fix_skills_section(text: str) -> str:
    """Find ## SKILLS, flatten its body, then rebuild one **Category:** per line."""
    m = re.search(
        r'(##\s*SKILLS[^\n]*\n)(.*?)(?=\n##\s|\Z)',
        text, re.DOTALL | re.IGNORECASE,
    )
    if not m:
        return text

    body = m.group(2)
    # Flatten all skill lines into one string so we can split on **Label:** boundaries
    # regardless of whether the model crammed them onto one line or spread them out.
    flat = ' '.join(ln.strip() for ln in body.split('\n') if ln.strip())

    pairs = re.findall(
        r'\*\*([^*]+)\*\*\s*:\s*(.*?)(?=\s*\*\*[^*]+\*\*\s*:|$)',
        flat,
    )
    if not pairs:
        return text

    rebuilt = []
    for label, content in pairs:
        label   = label.strip()
        content = re.sub(r'\s{2,}', ' ', content.strip().rstrip(',').strip())
        if label and content:
            rebuilt.append(f'**{label}:** {content}')

    new_body = '\n\n'.join(rebuilt) + '\n'
    return text[: m.start(2)] + new_body + text[m.end(2):]


def _remove_weak_content(text: str) -> str:
    """Remove blacklisted project entries that are too vague to add value."""
    bad = re.compile(r'Fintech\s+Client\s+Implementation', re.IGNORECASE)
    return '\n'.join(ln for ln in text.split('\n') if not bad.search(ln))


def _fix_experience_order(text: str) -> str:
    """Move freelance/consulting roles to the end of the EXPERIENCE section."""
    m = re.search(
        r'(##\s*(?:PROFESSIONAL\s+)?EXPERIENCE[^\n]*\n)(.*?)(?=\n##\s|\Z)',
        text, re.DOTALL | re.IGNORECASE,
    )
    if not m:
        return text

    body = m.group(2)
    # Split into individual role blocks — each starts with ###
    blocks = re.split(r'(?=^###)', body, flags=re.MULTILINE)
    blocks = [b for b in blocks if b.strip()]

    if len(blocks) <= 1:
        return text

    freelance_re = re.compile(r'freelan|consul|self.?employ', re.IGNORECASE)
    regular: list[str] = []
    freelance: list[str] = []
    for block in blocks:
        first_line = block.split('\n')[0]
        (freelance if freelance_re.search(first_line) else regular).append(block)

    if not freelance:
        return text  # nothing to reorder

    new_body = ''.join(regular + freelance)
    return text[: m.start(2)] + new_body + text[m.end(2):]


def _dedup_bullets(text: str) -> str:
    """Remove duplicate and empty bullet lines."""
    seen: set[str] = set()
    out: list[str] = []
    for line in text.split('\n'):
        s = line.strip()
        if s.startswith('- ') or s.startswith('* ') or s.startswith('• '):
            content = s[2:].strip()
            if not content or content in seen:
                continue
            seen.add(content)
        out.append(line)
    return '\n'.join(out)


def _clean_ai_output(text: str) -> str:
    """Full post-processing pipeline applied to every Groq response."""
    # 1. Strip code-block fences
    text = re.sub(r'```[\w]*\n?', '', text)
    text = re.sub(r'```', '', text)

    # 2. Discard everything before the first # heading (leaked preamble / analysis)
    name_m = re.search(r'(?m)^# ', text)
    if name_m and name_m.start() > 0:
        text = text[name_m.start():]

    # 3. Remove residual STEP headers and ━ separator bars that slipped through
    text = _strip_thinking_lines(text)

    # 4. Collapse triple/multiple colons  →  single colon
    text = re.sub(r':(\s*:)+', ':', text)

    # 5. Remove trailing colons at end of line
    text = re.sub(r':\s*$', '', text, flags=re.MULTILINE)

    # 6. Drop empty bullet lines
    text = re.sub(r'^[-*•]\s*$', '', text, flags=re.MULTILINE)

    # 7. Remove blacklisted vague project entries
    text = _remove_weak_content(text)

    # 8. Reconstruct the SKILLS section — guaranteed one category per line
    text = _fix_skills_section(text)

    # 9. Safety net: any **Category:** still inline gets its own line
    text = re.sub(
        r'([^\n])\s+(\*\*[A-Za-z][A-Za-z ,&/]+\*\*:)',
        r'\1\n\2',
        text,
    )

    # 10. Blank line between consecutive skill-category lines
    text = re.sub(
        r'(\*\*[A-Za-z][A-Za-z ,&/]+\*\*:[^\n]*)\n(\*\*[A-Za-z][A-Za-z ,&/]+\*\*:)',
        r'\1\n\n\2',
        text,
    )

    # 11. Move freelance/consulting roles to end of EXPERIENCE
    text = _fix_experience_order(text)

    # 12. Deduplicate and drop empty bullets
    text = _dedup_bullets(text)

    # 13. Collapse 3+ blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def extract_template_text(filename: str, content: bytes) -> str:
    """Extract plain text from a template file (pdf / docx / txt)."""
    ext = filename.rsplit('.', 1)[-1].lower()
    if ext == 'pdf':
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            return '\n'.join(
                page.extract_text() or '' for page in reader.pages
            ).strip()
        except Exception:
            return ''
    if ext == 'docx':
        try:
            doc = Document(io.BytesIO(content))
            return '\n'.join(p.text for p in doc.paragraphs).strip()
        except Exception:
            return ''
    if ext in ('txt', 'md'):
        return content.decode('utf-8', errors='ignore').strip()
    # images — we can't extract text without OCR; return empty
    return ''


# ── PDF renderer ───────────────────────────────────────────────────────────────

def _render_header_band(pdf: FPDF, L: int, name: str, subtitle: str) -> None:
    """Draw the dark navy header band with name + subtitle."""
    band_h = 30
    pdf.set_fill_color(*_NAVY)
    pdf.rect(0, 0, pdf.w, band_h, 'F')
    pdf.set_y(7)
    pdf.set_x(L)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font('Helvetica', 'B', 18)
    pdf.multi_cell(0, 8, name, align='L')
    pdf.set_x(L)
    pdf.set_font('Helvetica', '', 9)
    pdf.set_text_color(*_LIGHT)
    pdf.multi_cell(0, 5, subtitle, align='L')
    pdf.set_y(band_h + 4)
    pdf.set_text_color(*_DARK)


def markdown_to_pdf_bytes(md_text: str, candidate_name: str = "") -> bytes:
    """Professional resume / cover-letter PDF renderer."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=12)
    pdf.add_page()
    pdf.set_margins(0, 0, 0)

    L, R = 14, 14
    lines = [l.rstrip() for l in md_text.split('\n')]

    # ── Detect document type ────────────────────────────────────────────────
    first_content = next(
        (l.strip() for l in lines if l.strip() and not re.fullmatch(r'-{3,}', l.strip())),
        ''
    )
    is_cover = bool(re.search(r'COVER\s*LETTER', first_content, re.I))

    # ════════════════════════════════════════════════════════════════════════
    # COVER LETTER rendering
    # ════════════════════════════════════════════════════════════════════════
    if is_cover:
        display_name = candidate_name or "Applicant"
        _render_header_band(pdf, L, display_name, "Cover Letter")

        skip_title   = True   # skip the "# COVER LETTER" line (already in header)
        blank_streak = 0

        for raw in lines:
            s = raw.strip()

            if re.fullmatch(r'-{3,}', s):
                continue

            # Skip the # COVER LETTER heading — it's in the header band
            if skip_title and s.startswith('# '):
                skip_title = False
                continue

            if not s:
                if blank_streak == 0:
                    pdf.ln(4)          # generous paragraph gap for prose
                blank_streak += 1
                continue

            blank_streak = 0

            # Salutation: "Dear Hiring Manager,"
            if re.match(r'^Dear\b|^To\s+Whom', s, re.I):
                pdf.ln(2)
                pdf.set_x(L)
                pdf.set_font('Helvetica', 'B', 10)
                pdf.set_text_color(*_DARK)
                pdf.multi_cell(0, 5.5, _clean(s), align='L')
                pdf.ln(3)

            # Closing: "Sincerely," / "Regards," / "Best,"
            elif re.match(r'^(Sincerely|Regards|Best|Warm\s+regards|Yours)', s, re.I):
                pdf.ln(6)
                pdf.set_x(L)
                pdf.set_font('Helvetica', 'B', 10)
                pdf.set_text_color(*_DARK)
                pdf.multi_cell(0, 5.5, _clean(s), align='L')

            # Signature name (line immediately after closing)
            else:
                pdf.set_x(L)
                pdf.set_font('Helvetica', '', 10)
                pdf.set_text_color(*_MID)
                pdf.multi_cell(0, 5.5, _clean(s), align='L')

        return bytes(pdf.output())

    # ════════════════════════════════════════════════════════════════════════
    # RESUME rendering
    # ════════════════════════════════════════════════════════════════════════

    # First pass — pull name + contact line for the header band
    header_name    = None
    header_contact = None
    body_start     = 0

    for i, raw in enumerate(lines):
        s = raw.strip()
        if re.fullmatch(r'-{3,}', s):
            continue
        if s.startswith('# '):
            header_name = _clean(s[2:])
            for j in range(i + 1, min(i + 5, len(lines))):
                ns = lines[j].strip()
                if not ns or re.fullmatch(r'-{3,}', ns):
                    continue
                if ns.startswith('#'):
                    body_start = j
                    break
                header_contact = _clean(ns)
                body_start = j + 1
                break
            else:
                body_start = i + 1
            break
        else:
            body_start = 0
            break

    # Render header band
    if header_name:
        subtitle = header_contact or ""
        _render_header_band(pdf, L, header_name, subtitle)

    # Second pass — body
    blank_streak = 0

    for raw in lines[body_start:]:
        s = raw.strip()

        if re.fullmatch(r'-{3,}', s):
            continue

        # h2 — section heading with accent bar
        if s.startswith('## '):
            blank_streak = 0
            pdf.ln(4)
            bar_y = pdf.get_y()
            pdf.set_fill_color(*_ACCENT)
            pdf.rect(L, bar_y, 3, 5.5, 'F')
            pdf.set_x(L + 5)
            pdf.set_font('Helvetica', 'B', 10)
            pdf.set_text_color(*_ACCENT)
            pdf.multi_cell(0, 5.5, _clean(s[3:]).upper(), align='L')
            y = pdf.get_y()
            pdf.set_draw_color(*_RULE)
            pdf.line(L, y, pdf.w - R, y)
            pdf.ln(2)
            pdf.set_text_color(*_DARK)

        # h3 — job title / degree
        elif s.startswith('### '):
            blank_streak = 0
            pdf.ln(3)
            pdf.set_x(L)
            pdf.set_font('Helvetica', 'B', 9)
            pdf.set_text_color(*_DARK)
            pdf.multi_cell(0, 4.8, _clean(s[4:]), align='L')

        # bullet
        elif s.startswith('- ') or s.startswith('* '):
            blank_streak = 0
            pdf.set_x(L + 5)
            pdf.set_font('Helvetica', '', 9)
            pdf.set_text_color(*_MID)
            pdf.multi_cell(0, 4.5, '- ' + _clean(s[2:]), align='L')

        # blank line
        elif not s:
            if blank_streak == 0:
                pdf.ln(1.5)
            blank_streak += 1
            continue

        # paragraph / contact / inline text
        else:
            blank_streak = 0
            pdf.set_x(L)
            pdf.set_font('Helvetica', '', 9)
            pdf.set_text_color(*_MID)
            pdf.multi_cell(0, 4.5, _clean(s), align='L')

        blank_streak = 0

    return bytes(pdf.output())


# ── DOCX renderer ──────────────────────────────────────────────────────────────

def _clean_docx(text: str) -> str:
    """Strip markdown markers for DOCX — no latin-1 encode, python-docx handles UTF-8."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',     r'\1', text)
    text = re.sub(r'__(.+?)__',     r'\1', text)
    text = re.sub(r'`(.+?)`',       r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]',  r'\1', text)
    text = text.translate(_UNICODE_MAP)
    return text.strip()


def _add_paragraph_border(para) -> None:
    """Add an indigo bottom border to a paragraph (used for section headings)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '3730A3')
    pBdr.append(bottom)
    pPr.append(pBdr)


def markdown_to_docx_bytes(md_text: str, candidate_name: str = "") -> bytes:
    """Convert markdown resume to a clean Word (.docx) document."""
    doc = Document()

    # Page margins
    for sec in doc.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(2.2)

    # Remove the default empty paragraph python-docx adds
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    lines = md_text.split('\n')
    first_h1_done = False

    for raw in lines:
        s = raw.strip()

        if re.fullmatch(r'-{3,}', s):
            continue

        if not s:
            # Minimal spacing between blocks
            doc.add_paragraph()
            continue

        if s.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(_clean_docx(s[2:]))
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(17, 24, 39)
            first_h1_done = True

        elif s.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(_clean_docx(s[3:]).upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(55, 48, 163)
            _add_paragraph_border(p)

        elif s.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(_clean_docx(s[4:]))
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(17, 24, 39)

        elif s.startswith('- ') or s.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(_clean_docx(s[2:]))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(55, 65, 81)

        else:
            p = doc.add_paragraph()
            # Contact / subtitle line (line right after the name)
            if first_h1_done and not any(
                ln.strip().startswith('## ')
                for ln in lines[: lines.index(raw)]
                if ln.strip()
            ):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(_clean_docx(s))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(55, 65, 81)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── routes ─────────────────────────────────────────────────────────────────────

@app.get("/")
def root():
    return {"status": "AI Resume Generator API running"}


@app.post("/generate-resume")
async def generate_resume(
    job_title: str        = Form(...),
    job_description: str  = Form(...),
    existing_resume: str  = Form(""),
    name: str             = Form(""),
    email: str            = Form(""),
    phone: str            = Form(""),
    location: str         = Form(""),
    experience: str       = Form(""),
    skills: str           = Form(""),
    education: str        = Form(""),
    custom_instructions: str = Form(""),
    template_file: UploadFile = File(None),
):
    def val(field: str, fallback: str) -> str:
        return field.strip() if field.strip() else fallback

    has_existing = existing_resume.strip() != ""

    # ── candidate block ──────────────────────────────────────────────────────
    loc = val(location, 'Extract from existing resume') if has_existing else location.strip()

    # Trim large user-supplied fields so the prompt stays within model token limits.
    # Groq's smallest fallback (llama-3.1-8b-instant) allows ~6 000 TPM;
    # the fixed prompt template itself is ~2 500 tokens, leaving ~3 000 for data.
    _jd   = job_description[:1500]
    _er   = existing_resume[:2500]
    _exp  = experience[:800]
    _sk   = skills[:400]
    _edu  = education[:400]

    if has_existing:
        candidate_section = f"""CANDIDATE DETAILS (extract any missing fields from EXISTING RESUME below):
- Name: {val(name, 'Extract from existing resume')}
- Email: {val(email, 'Extract from existing resume')}
- Phone: {val(phone, 'Extract from existing resume')}
- Location: {val(location, 'Extract from existing resume')}
- Target Job: {job_title}
- Job Description: {_jd}
- Experience: {val(experience, 'Extract from existing resume')}
- Skills: {val(skills, 'Extract from existing resume')}
- Education: {val(education, 'Extract from existing resume')}

EXISTING RESUME — use as the base, preserve all real data, tailor to target job:
{_er}"""
    else:
        candidate_section = f"""CANDIDATE DETAILS:
- Name: {name}
- Email: {email}
- Phone: {phone}
- Location: {location}
- Target Job: {job_title}
- Job Description: {_jd}
- Experience: {_exp}
- Skills: {_sk}
- Education: {_edu}"""

    # ── template block ───────────────────────────────────────────────────────
    has_template    = False
    tpl_has_summary = True   # default: include summary when no template
    template_section = ""

    if template_file and template_file.filename:
        raw_bytes = await template_file.read()
        tpl_text  = extract_template_text(template_file.filename, raw_bytes)
        if tpl_text:
            has_template    = True
            tpl_has_summary = bool(re.search(
                r'\bSUMMARY\b|\bPROFILE\b|\bOBJECTIVE\b|\bABOUT\b',
                tpl_text.upper()
            ))
            _c_note = (
                "C. NO SUMMARY — the template contains no summary/profile/objective section; "
                "DO NOT add one under any circumstances"
                if not tpl_has_summary else
                "C. SUMMARY — the template includes a summary section; reproduce it"
            )
            template_section = f"""

╔════════════════════════════════════════════════════════════╗
║  CRITICAL: RESUME TEMPLATE PROVIDED — OVERRIDE ALL DEFAULTS ║
╚════════════════════════════════════════════════════════════╝
The user uploaded a resume template. You MUST mirror its structure EXACTLY.
All section defaults in STEP 4 are OVERRIDDEN by this template.

TEMPLATE CONTENT (read carefully — use its section order and names):
---
{tpl_text[:1500]}
---

MANDATORY TEMPLATE RULES (highest priority — override everything else):
A. SECTION ORDER — reproduce sections in the exact order they appear in the template above
B. SECTION NAMES — use the exact heading text from the template (e.g. "WORK HISTORY" not "PROFESSIONAL EXPERIENCE")
{_c_note}
D. SKILLS FORMAT — if the template shows skills as separate labeled lines, keep one category per line; never merge into a single paragraph
E. BULLETS — every experience and project entry must use bullet points; NEVER collapse bullets into prose paragraphs
F. ALL PROJECTS — include every project from the candidate data; never truncate or stop at 2 or 3
G. ALL CERTS — include every certification from the candidate data; never truncate
H. NO EXTRAS — do not insert any section that does not appear in the template"""
        else:
            template_section = "\n\nTEMPLATE NOTE: A visual template was uploaded. Use a clean professional layout with a coloured header."

    # ── hints ────────────────────────────────────────────────────────────────
    contact_hint = (
        "[Email] | [Phone] | [Location]"
        if has_existing and not email.strip()
        else f"{email} | {phone}{' | ' + loc if loc else ''}"
    )
    name_hint = (
        "[Name from resume]" if has_existing and not name.strip() else name
    )

    custom_block = (
        f"\nCANDIDATE CUSTOM INSTRUCTIONS (highest priority — override defaults where needed):\n{custom_instructions.strip()}"
        if custom_instructions.strip() else ""
    )

    # ── build STEP 4 block (different in template mode vs default mode) ──────
    if has_template:
        _summary_step = (
            "- DO NOT include a Professional Summary — the template has none"
            if not tpl_has_summary else
            "- Include a Professional Summary (follow the template's format)"
        )
        step4_block = f"""━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
STEP 4 — WRITE THE RESUME  [TEMPLATE MODE — ignore default structure]
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
A template was provided above. DO NOT use the default section structure.
Instead, look at the TEMPLATE CONTENT above and reproduce it exactly.

Begin with:
# [Candidate Full Name]
[Email] | [Phone] | [Location] | LinkedIn: [URL if available]

Then reproduce every section from the template in order, filling with candidate data:
{_summary_step}
- All experience/project items must be bullet points (- item), never prose
- Skills: one labeled category per line, never merged into one block
- Projects: include ALL from the candidate data — never stop at 2 or 3
- Certifications: include ALL — never truncate the list

After the resume body, append the cover letter."""
    else:
        step4_block = f"""━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
STEP 4 — WRITE THE RESUME
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Output in this exact markdown structure:

---
# {name_hint}
{contact_hint} | LinkedIn: [URL if available]

---
## PROFESSIONAL SUMMARY
[2-3 punchy sentences. Open with seniority + domain. Name the target role. Close with unique value proposition. Tone must match the company type detected in Step 2.]

---
## SKILLS
**Technical:** skill1, skill2, skill3

**Tools & Platforms:** tool1, tool2, tool3

[Add further categories only if genuinely applicable — one per line, blank line between each]

SKILLS FORMAT LAW: each category = its own line with a blank line above and below it. "**CategoryName:** item1, item2, item3". NEVER run two categories onto the same line. NEVER put all skills in one unlabelled block.

---
## PROFESSIONAL EXPERIENCE
[Roles STRICTLY by start date — MOST RECENT FIRST. Exception: freelance / consulting / self-employed roles that overlap with other employment must be listed AFTER all non-freelance roles, regardless of their dates. 4-5 bullets for the most recent/relevant role, 2-3 for older ones.]
### [Exact Job Title] | [Company] | [Start Month Year] - [End Month Year or Present]
- [Achievement bullet with metric]
- [Achievement bullet with metric]
- [Achievement bullet with metric]

---
## PROJECTS
[Include ONLY if the candidate data mentions projects. Omit this section entirely if no projects are present.]
### [Project Name] | [Tech stack or context]
- [What it does / your role / impact]

---
## EDUCATION
### [Degree] | [Institution] | [Year]

---
## CERTIFICATIONS & ACHIEVEMENTS
- [Only include if genuinely present in source data]

---"""

    # ── template-specific hard rules (appended to main rules) ───────────────
    if has_template:
        _no_sum_rule = (
            "19. TEMPLATE MODE — DO NOT include a PROFESSIONAL SUMMARY section; the template has none"
            if not tpl_has_summary else
            "19. TEMPLATE MODE — include PROFESSIONAL SUMMARY as shown in the template"
        )
        template_hard_rules = f"""
22. TEMPLATE MODE — section order and section names must EXACTLY match the uploaded template; do not reorder or rename
{_no_sum_rule.replace("19.", "23.")}
24. TEMPLATE MODE — include EVERY project and EVERY certification from the candidate data; never truncate"""
    else:
        template_hard_rules = ""

    prompt = f"""You are a world-class resume strategist and ATS optimization expert. Your job is to think and write like a senior executive recruiter — not fill in a template.

⚠ IMPORTANT — READ BEFORE ANYTHING ELSE:
Your internal analysis in Steps 1-3 is for your thinking only and must NEVER appear in your output.
NEVER output Step 1, Step 2, or Step 3 headings or analysis text.
Output ONLY the formatted resume followed by the cover letter.
The very first line of your response must be the candidate's name in the format: # First Last

{candidate_section}{template_section}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
INTERNAL STEP 1 — ANALYZE CANDIDATE PROFILE  [THINKING ONLY — DO NOT OUTPUT]
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Before writing anything, internally assess:
- Years of total experience → determines one-page (0-7 yrs) vs two-page (8+ yrs) resume
- Seniority level → entry / mid / senior / lead / executive
- Industry domain → tech, finance, healthcare, product, data, etc.
- Biggest 3 strengths most relevant to the target role
- Any gaps, career pivots, or unusual patterns to handle carefully

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
INTERNAL STEP 2 — ANALYZE TARGET ROLE  [THINKING ONLY — DO NOT OUTPUT]
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
From the job title and description, identify:
- Top 5-8 keywords and skills the ATS will scan for → weave every one naturally into bullets and summary
- Tone of the company: startup (energetic, impact-driven) vs enterprise (professional, process-focused) vs agency (creative, client-oriented) → match that tone throughout
- Whether the role emphasises individual contribution, leadership, or cross-functional collaboration → lead with that angle

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
INTERNAL STEP 3 — TRANSFORM BULLETS  [THINKING ONLY — DO NOT OUTPUT]
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
For every bullet point:
- Start with a powerful action verb (Architected, Spearheaded, Reduced, Scaled, Automated, Delivered…)
- Add a metric or result: %, $, time, scale, users, speed, ranking, team size
- If the source data is vague (e.g. "worked on APIs"), intelligently infer a realistic metric based on context — but NEVER invent company names, titles, or dates
- Connect the achievement to a business outcome when possible

{step4_block}

---
# COVER LETTER

Dear Hiring Manager,

[Opening — 2-3 sentences: name the exact role and company, state your strongest relevant qualification, show genuine excitement with a specific reason why this company/role appeals to you]

[Body paragraph 1 — specific technical or project achievement with metrics that directly maps to a key requirement from the JD]

[Body paragraph 2 — leadership, collaboration, or soft-skill achievement that addresses another JD requirement; mention company culture fit]

[Closing — 1-2 sentences: express enthusiasm, include a clear call to action, mention availability]

Sincerely,
{name_hint}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
HARD RULES — NEVER BREAK THESE
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. NEVER output placeholder text like [Company Name] or [Add here] — use real data or omit the field
2. NEVER invent or hallucinate company names, job titles, or dates not present in the source
3. NEVER use code blocks, backticks, or ```markdown markers — plain markdown only
4. SKILLS: every category on its own line — format is exactly "**Category:** item1, item2, item3" — NEVER merge two categories onto the same line, NEVER dump all skills into one unlabelled block
5. Use only hyphens (-) for date ranges — never Unicode dashes (–, —)
6. Single blank line between sections only — never double blank lines
7. Every bullet must start with a strong past-tense action verb (present tense for current role)
8. Cover letter must reference the exact job title: {job_title}
9. Skills: NLP, Computer Vision, Machine Learning, AI are TECHNIQUES → **Technical:**, not **Tools & Platforms:**
10. NEVER mention a company name in the Professional Summary — reference the role title only, never "at [Company]"
11. NEVER invent metrics — if no number exists in the source data, use a strong action verb without a fabricated percentage or figure
12. ALWAYS include a PROJECTS section between EXPERIENCE and EDUCATION when the candidate data mentions any projects; omit the section entirely if no projects are present
13. SKILLS: **Technical:** and **Tools & Platforms:** must each be on their own line with a blank line between every category pair — never adjacent with no gap
14. NEVER output empty bullet points — if a bullet has no real content, omit it entirely; never write "- " with nothing after it
15. NEVER duplicate bullet points — if the same sentence or achievement appears more than once in the resume, include it exactly once
16. EXPERIENCE: sort roles strictly by start date, most recent first — freelance / consulting / self-employed roles that overlap with other employment must appear AFTER all regular employment roles
17. NEVER output multiple colons in a row (e.g. "Label: : :" is forbidden) — use a single colon; NEVER leave a trailing colon with nothing after it
18. NEVER add sub-bullets or expanded bullets under project descriptions — each project entry is ONE header line (### Project Name) followed by a single description bullet maximum; never expand or add more bullets
19. NEVER output STEP 1, STEP 2, STEP 3, or any analytical thinking text — output the resume and cover letter only; the very first line must be "# CandidateName"
20. NEVER include a project named "Fintech Client Implementation" — omit it entirely; it is too vague to add value
21. SKILLS: **Technical:** and **Tools & Platforms:** must ALWAYS be on completely separate lines with a blank line between them — this takes priority over everything else
{template_hard_rules}
{custom_block}
"""

    cleaned = _clean_ai_output(_chat(prompt, temperature=0.5))
    return {"resume": cleaned, "status": "success"}


@app.post("/refine-resume")
async def refine_resume(
    current_resume: str = Form(...),
    instruction: str    = Form(...),
):
    prompt = f"""You are an expert resume editor. Apply the requested change to the resume below.

CURRENT RESUME & COVER LETTER:
{current_resume}

REQUESTED CHANGE:
{instruction}

RULES:
1. Apply ONLY the requested change — do not alter anything else
2. Return the COMPLETE updated document in the same markdown format
3. NEVER use code blocks, backticks, or ```markdown markers
4. Keep all section headings, structure, and order identical
5. Preserve all real data — never invent or remove facts
6. SKILLS: each category must remain on its own line — "**Category:** item1, item2, item3" — never merge two categories onto the same line
7. NEVER output empty bullets, duplicate bullets, multiple consecutive colons, or trailing colons with nothing after them
"""

    cleaned = _clean_ai_output(_chat(prompt, temperature=0.3))
    return {"resume": cleaned, "status": "success"}


@app.post("/score-resume")
async def score_resume(
    resume_text: str     = Form(...),
    job_description: str = Form(""),
):
    jd_block = (
        f"Job Description to match against:\n{job_description[:1500]}"
        if job_description.strip()
        else "No job description provided — score against general ATS best practices for the role."
    )

    prompt = f"""You are an ATS (Applicant Tracking System) expert and resume evaluator.

{jd_block}

RESUME TO SCORE:
{resume_text[:3000]}

Return ONLY a valid JSON object — no markdown, no explanation, no code fences, nothing else:

{{
  "overall_score": <integer 0-100>,
  "keyword_match": <integer 0-100>,
  "formatting_score": <integer 0-100>,
  "impact_score": <integer 0-100>,
  "length_score": <integer 0-100>,
  "improvements": [
    "<specific, actionable suggestion referencing actual resume content>",
    "<specific, actionable suggestion referencing actual resume content>",
    "<specific, actionable suggestion referencing actual resume content>"
  ],
  "strengths": [
    "<specific strength found in this resume>",
    "<specific strength found in this resume>",
    "<specific strength found in this resume>"
  ]
}}

Scoring criteria:
- overall_score: weighted average (keyword 30% + impact 30% + formatting 20% + length 20%)
- keyword_match: presence of role-relevant and JD keywords (or general industry keywords if no JD)
- formatting_score: clear sections, consistent structure, ATS-parseable hierarchy
- impact_score: quantified achievements, action verbs, specific metrics and results
- length_score: appropriate for experience (one page ideal for under 7 years; two pages fine for 7+)
- Be honest and critical. A score of 100 is nearly impossible. A strong resume scores 68-80.
- improvements and strengths must be SPECIFIC to THIS resume — no generic filler advice.
"""

    raw = _chat(prompt, temperature=0.1)

    # Extract JSON even if model wraps it in prose
    match = re.search(r'\{[\s\S]*\}', raw)
    if not match:
        raise HTTPException(status_code=500, detail="Scorer returned no JSON")

    try:
        data = json.loads(match.group())
    except json.JSONDecodeError as e:
        raise HTTPException(status_code=500, detail=f"JSON parse error: {e}")

    # Clamp all numeric scores to 0-100
    for key in ("overall_score", "keyword_match", "formatting_score", "impact_score", "length_score"):
        if key in data:
            data[key] = max(0, min(100, int(data[key])))

    return data


@app.post("/download-pdf")
async def download_pdf(
    text: str    = Form(...),
    section: str = Form("resume"),
    name: str    = Form("Candidate"),
):
    try:
        resume_text, cover_text = split_content(text)

        if section == "cover_letter":
            content  = cover_text if cover_text else text
            filename = f"{name.replace(' ', '_')}_Cover_Letter.pdf"
        else:
            content  = resume_text if resume_text else text
            filename = f"{name.replace(' ', '_')}_Resume.pdf"

        pdf_bytes = markdown_to_pdf_bytes(content, candidate_name=name)

        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {e}")


@app.post("/download-docx")
async def download_docx_endpoint(
    text: str = Form(...),
    name: str = Form("Candidate"),
):
    try:
        resume_text, _ = split_content(text)
        content  = resume_text if resume_text else text
        filename = f"{name.replace(' ', '_')}_Resume.docx"

        docx_bytes = markdown_to_docx_bytes(content, candidate_name=name)

        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {e}")
